import json
import docx
import sys
import docx
import regex as re
from typing import List, Dict
from pydantic import BaseModel
from openai import AzureOpenAI
import os
from dotenv import load_dotenv
import docx
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import uuid

#from docx import Document
from docx import Document  # Use bayoo-docx for native comment support
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

load_dotenv()

def load_markdown_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        return content
    except FileNotFoundError:
        return "The file was not found."
    except Exception as e:
        return f"An error occurred: {e}"
    

def extract_text_by_page(filepath):
    """
    Opens a .docx file and extracts its contents as a list of strings, where each entry represents a page.
    
    Parameters:
        filepath (str): Path to the .docx file.
    
    Returns:
        list: List of strings, each corresponding to a page in the document.
    """
    doc = docx.Document(filepath)
    pages = []
    current_page = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # Ignore empty lines
            current_page.append(text)
        
        # Here, we assume that a page break is denoted by a specific paragraph style or a form of delimiter.
        # If explicit page breaks exist in the document, they should be handled here.
    
    # Since python-docx doesn't support page-level extraction, return a single text blob.
    # Consider a more advanced library like `pdfplumber` or `PyMuPDF` for actual page-based parsing.
    return ["\n".join(current_page)]  # Returns entire content as a single 'page' unless improved


def find_hyphenated_terms(text):
    """
    Uses regex to find words linked by a hyphen or dash (e.g., carbon-neutral, 2010–2014),
    while excluding cases where the dash is used as a sentence break.

    Parameters:
        text (str): The input text to search through.

    Returns:
        list: A list of dictionaries, each containing:
            - 'match': The matched hyphenated term.
            - 'start': The starting index of the match.
            - 'end': The ending index of the match.
    """
    pattern = r'\b[\p{L}\p{N}]+\s*[-–—]\s*[\p{L}\p{N}]+\b'
    
    matches = []
    for match in re.finditer(pattern, text):
        # Extract match details
        match_text = match.group()
        start, end = match.span()

        # Exclude sentence-break hyphens/dashes using a naive heuristic:
        # If the match is surrounded by spaces and contains dashes with spaces, it's likely a sentence break
        if not re.search(r'\s[-–—]\s', match_text):  
            matches.append({'match_term': match_text, 'start': start, 'end': end})
    
    return matches


def unique_hyphenated_terms(matches):
    """
    Given a list of hyphenated terms, return a list of unique hyphenated terms.
    """
    unique_matches = {}
    for match_term in set([match['match_term'] for match in matches]):
        for match in matches:
            if match_term == match['match_term']:
                unique_matches[match_term] = match
                break
    return unique_matches

def add_context_to_matches(text, matches):
    """
    Adds context sentences to each match in the list of matches.

    Parameters:
        text (str): The full text where matches were found.
        matches (list): A list of dictionaries containing the matches from the previous function.

    Returns:
        list: The updated list of matches with an added 'context_sentence' key.
    """
    sentences = re.split(r'(?<=[.!?])\s+', text)  # Split text into sentences
    
    for match in matches:
        match_start = match['start']
        
        # Find the sentence containing the match
        for sentence in sentences:
            if match_start in range(text.find(sentence), text.find(sentence) + len(sentence)):
                match['context_sentence'] = sentence
                break
    
    return matches


# Define the structured output model
class ValidationResult(BaseModel):
    case_id: int
    match_term: str
    adheres_to_rule: bool

class CasesList(BaseModel):
    cases: List[ValidationResult]

# Initialize Azure OpenAI client
client = AzureOpenAI(
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version="2024-02-01",
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
)


def get_azure_client() -> AzureOpenAI:
    """Initialize Azure OpenAI client"""
    print("Initializing Azure OpenAI client...")
    try:
        client = AzureOpenAI(
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
        )
        print("Azure OpenAI client initialized successfully")
        return client
    except Exception as e:
        print(f"Error initializing Azure OpenAI client: {str(e)}")
        raise
    
    

def check_adherence_to_guidelines(lexical_guideline: str, cases: List[Dict], client=None) -> List[Dict]:
    """
    Check if each case adheres to the lexical guideline using Azure OpenAI.

    Args:
        lexical_guideline (str): The lexical guideline in markdown format.
        cases (List[Dict]): A list of cases to be checked.

    Returns:
        List[Dict]: The list of cases with adherence results.
    """
    if client is None:
        client = get_azure_client()
        
    # Prepare System Prompt
    system_prompt = f"You are proofreader at the European Commission who is responsible for checking adherence of sententences to the lexical rules of the typographical styleguide.\n"
    system_prompt += f"You will be given a list of cases, each with a match and a context sentence. You will need to check if the match adheres to the lexical rules.\n"
    
    # Prepare the prompt
    prompt = f"Lexical Guideline in Markdown format:\n\n```hyphenation_rule.md\n{lexical_guideline}\n```\n\n"
    prompt += "Please validate the following cases:\n\n```cases.txt\n"
    for i, case in enumerate(cases, 1):
        prompt += f"Case_id: {i}\nMatch_term: {case['match_term']}\nContext: {case['context_sentence']}\n{'-'*50}\n"
    prompt += "```\n For each case, read the case, read the lexical rule, identify which rule is applicable, and then make your judgement.\n"

    # Call the Azure OpenAI model
    response = client.beta.chat.completions.parse(
        model=os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME"),  # Replace with your deployed model name
        messages=[{"role": "system", "content": system_prompt},
                  {"role": "user", "content": prompt}],
        response_format=CasesList,
        temperature=0.1,
        #strict=True  # Enforce adherence to the response model
    )
    #print(f"\n\nResponse: {response}\n\n\n")
    #print(f"\nResponse Content: {response.choices[0].message.parsed}\n\n\n")
    
    adherence_cases_json = json.loads(response.choices[0].message.parsed.model_dump_json(indent=2))['cases']
    adherence_cases_dict = {result_case['match_term']: result_case['adheres_to_rule'] for result_case in adherence_cases_json}
    return adherence_cases_dict
    # Parse the response and update cases
    # for i, result in enumerate(response):
        # cases[i]['adheres_to_rule'] = result.adheres_to_rule
# 
    # return cases


def append_adherence_to_cases(cases_list: List[Dict], adherence_cases_dict: Dict, final_cases: List[Dict]=None) -> List[Dict]:
    if final_cases is None:
        final_cases = []
    for case_item in cases_list:
        final_item = case_item.copy()
        final_item['adheres_to_rule'] = adherence_cases_dict[case_item['match_term']]
        print(f"Appending case: {final_item}")
        final_cases.append(final_item)
    return final_cases


def add_comments_to_docx(docx_path, json_path, output_path=None):
    """
    Adds comments to a .docx file for terms marked as 'adheres_to_rule': False in the results.json file.

    Parameters:
        docx_path (str): Path to the .docx file.
        json_path (str): Path to the JSON results file.

    Output:
        Saves a new `.docx` file with comments added for incorrect hyphenations.
    """
    # Load the docx file using bayoo-docx
    doc = Document(docx_path)

    # Load results JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        results = json.load(f)

    # Filter out cases that do not adhere to the rule
    incorrect_cases = [case for case in results if not case["adheres_to_rule"]]

    # Unique ID generator for comments
    comment_id_counter = 1

    # Process paragraphs to find exact matches
    for case in incorrect_cases:
        match_text = case["match_term"]

        for para in doc.paragraphs:
            if match_text in para.text:
                found = False
                for run in para.runs:
                    if match_text in run.text:
                        # Insert comment using bayoo-docx's built-in comment support
                        para.add_comment(
                            text=f"Incorrect hyphenation: '{match_text}' should be revised.",
                            author="Proofreader",
                            
                        )
                        comment_id_counter += 1
                        found = True
                        break
                if found:
                    break  # Move to the next case once matched

    # Generate output filename
    if output_path is None:
        output_path = generate_output_filename(docx_path)

    # Save the modified document
    doc.save(output_path)
    print(f"✅ Comments added and document saved to {output_path}")

def generate_output_filename(docx_path):
    """
    Generates a new filename with `_hyphenation_comments` before the `.docx` extension.

    Parameters:
        docx_path (str): The original file path.

    Returns:
        str: The modified output file path.
    """
    uuid_str = str(uuid.uuid1()).split('-')[0]
    base, ext = os.path.splitext(docx_path)
    if 'hyphenation' in base.split('/')[-1]:
        return f"{base}_comments_{uuid_str}{ext}"
    else:
        return f"{base}_hyphenation_comments_{uuid_str}{ext}"